"""根据 test2.py 提取的键值对，生成与参考文档表格一致的 test.docx。"""
import sys
sys.path.insert(0, r'c:\Users\swm\Desktop\wordtoexcel')
from test2 import extract_table_kv
from docx import Document
from docx.shared import Pt, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def _set_cell_border(cell, **kwargs):
    """设置单元格边框。"""
    tc = cell._tc
    tcpr = tc.get_or_add_tcPr()
    borders = tcpr.find(qn('w:tcBorders'))
    if borders is None:
        borders = OxmlElement('w:tcBorders')
        tcpr.insert(0, borders)
    for edge in ('top', 'bottom', 'left', 'right'):
        if edge in kwargs:
            attr = kwargs[edge]
            el = OxmlElement(f'w:{edge}')
            el.set(qn('w:val'), attr.get('val', 'single'))
            el.set(qn('w:sz'), attr.get('sz', '6'))
            el.set(qn('w:space'), attr.get('space', '0'))
            el.set(qn('w:color'), attr.get('color', 'auto'))
            borders.append(el)


def _set_cell_shading(cell, fill='FFFFFF'):
    tc = cell._tc
    tcpr = tc.get_or_add_tcPr()
    shd = tcpr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tcpr.append(shd)
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill)


def _set_cell_v_align(cell, val='center'):
    tc = cell._tc
    tcpr = tc.get_or_add_tcPr()
    va = tcpr.find(qn('w:vAlign'))
    if va is None:
        va = OxmlElement('w:vAlign')
        tcpr.append(va)
    va.set(qn('w:val'), val)


def _set_tbl_grid(table, widths_dxa):
    """设置表格列宽网格。"""
    tbl = table._tbl
    tblGrid = tbl.find(qn('w:tblGrid'))
    if tblGrid is not None:
        tbl.remove(tblGrid)
    tblGrid = OxmlElement('w:tblGrid')
    for w in widths_dxa:
        gc = OxmlElement('w:gridCol')
        gc.set(qn('w:w'), str(w))
        tblGrid.append(gc)
    tbl.insert(1, tblGrid)  # after tblPr


def _set_cell_width(cell, width_dxa):
    """设置单元格宽度 (dxa)。"""
    tcpr = cell._tc.get_or_add_tcPr()
    tcw = tcpr.find(qn('w:tcW'))
    if tcw is None:
        tcw = OxmlElement('w:tcW')
        tcpr.append(tcw)
    tcw.set(qn('w:w'), str(width_dxa))
    tcw.set(qn('w:type'), 'dxa')


def _style_cell(cell, top_sz='12', bottom_sz='6', right_sz='6', width_dxa=None):
    """统一单元格样式。"""
    _set_cell_border(cell,
                     top={'val': 'single', 'sz': top_sz, 'space': '0', 'color': 'auto'},
                     bottom={'val': 'single', 'sz': bottom_sz, 'space': '0', 'color': 'auto'},
                     right={'val': 'single', 'sz': right_sz, 'space': '0', 'color': 'auto'})
    _set_cell_shading(cell, 'FFFFFF')
    _set_cell_v_align(cell, 'center')
    if width_dxa is not None:
        _set_cell_width(cell, width_dxa)
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if not p.runs:
            p.add_run('')
        for r in p.runs:
            r.font.size = Pt(10.5)


def _lookup(kv: dict, key) -> str:
    """从键值对查找，兼容中文/英文冒号差异。"""
    if not key:
        return ''
    if key in kv:
        return kv[key]
    # 尝试替换冒号
    for alt in [key.replace('：', ':'), key.replace(':', '：')]:
        if alt in kv:
            return kv[alt]
    return ''


def generate_docx(kv: dict, output_path: str):
    doc = Document()

    # --- 默认字体 ---
    style = doc.styles['Normal']
    style.font.size = Pt(10.5)

    # --- 页面设置 ---
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

    # --- 标题 ---
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_title.add_run('《程序设计基础》课程质量评价报告')
    r.bold = True
    r.font.size = Pt(15)

    # --- 空行 ---
    doc.add_paragraph()

    # --- 副标题 ---
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_sub.add_run('课程基本信息表').font.size = Pt(10.5)

    # --- 列宽 (dxa) ---
    # 与参考文档一致：col0=1353, col1=3757, col2=1493, col3=1990
    widths = [1353, 3757, 1493, 1990]
    merged_w = widths[1] + widths[2] + widths[3]  # 7240

    # --- 字段映射 (左标签, 左值key, 右标签, 右值key, 是否整行合并) ---
    fields = [
        ('课程名称：', '课程名称：',    '修课人数：', None,          False),
        ('课程编号：', '课程编号：',    '课序号：',   None,          False),
        ('授课班级：', '适用专业：',    '学时数：',   '学时：总学时', False),
        ('选用教材：', None,            '学分：',     '学分：',       False),
        ('开课院系',   '开课单位:',     '授课教师：', '大纲执笔人：', False),
        ('课程性质：', '课程性质：',    None,         None,          True),
        ('课程类型：', '课程类型：',    None,         None,          True),
    ]

    # --- 创建表格 ---
    table = doc.add_table(rows=len(fields), cols=4)
    table.style = 'Table Grid'
    _set_tbl_grid(table, widths)

    for r_idx, (l_label, l_key, r_label, r_key, full_merge) in enumerate(fields):
        if full_merge:
            # 标签列（col 0）
            c0 = table.cell(r_idx, 0)
            c0.text = ''
            _style_cell(c0, top_sz='6', bottom_sz='12', width_dxa=widths[0])
            p0 = c0.paragraphs[0]
            p0.add_run(l_label).font.size = Pt(10.5)

            # 合并列 1-3
            c1 = table.cell(r_idx, 1)
            c2 = table.cell(r_idx, 2)
            c3 = table.cell(r_idx, 3)
            merged = c1.merge(c3)
            tcpr = merged._tc.get_or_add_tcPr()
            tcw = OxmlElement('w:tcW')
            tcw.set(qn('w:w'), str(merged_w))
            tcw.set(qn('w:type'), 'dxa')
            old_tcw = tcpr.find(qn('w:tcW'))
            if old_tcw is not None:
                tcpr.remove(old_tcw)
            tcpr.append(tcw)
            gs = tcpr.find(qn('w:gridSpan'))
            if gs is None:
                gs = OxmlElement('w:gridSpan')
                tcpr.append(gs)
            gs.set(qn('w:val'), '3')

            merged.text = ''
            _style_cell(merged, top_sz='6', bottom_sz='12', right_sz='6', width_dxa=merged_w)
            pv = merged.paragraphs[0]
            value_text = _lookup(kv, l_key) if l_key else ''
            pv.add_run(value_text).font.size = Pt(10.5)
        else:
            bt = '12' if r_idx == len(fields) - 1 else '6'
            tt = '12' if r_idx == 0 else '6'
            # 左标签
            c0 = table.cell(r_idx, 0)
            c0.text = ''
            _style_cell(c0, top_sz=tt, bottom_sz=bt, width_dxa=widths[0])
            c0.paragraphs[0].add_run(l_label).font.size = Pt(10.5)

            # 左值
            c1 = table.cell(r_idx, 1)
            c1.text = ''
            _style_cell(c1, top_sz=tt, bottom_sz=bt, width_dxa=widths[1])
            c1.paragraphs[0].add_run(_lookup(kv, l_key) if l_key else '').font.size = Pt(10.5)

            # 右标签
            c2 = table.cell(r_idx, 2)
            c2.text = ''
            _style_cell(c2, top_sz=tt, bottom_sz=bt, width_dxa=widths[2])
            c2.paragraphs[0].add_run(r_label or '').font.size = Pt(10.5)

            # 右值
            c3 = table.cell(r_idx, 3)
            c3.text = ''
            _style_cell(c3, top_sz=tt, bottom_sz=bt, width_dxa=widths[3])
            c3.paragraphs[0].add_run(_lookup(kv, r_key) if r_key else '').font.size = Pt(10.5)

    # 整表外框加粗
    tbl = table._tbl
    tblpr = tbl.find(qn('w:tblPr'))
    if tblpr is None:
        tblpr = OxmlElement('w:tblPr')
        tbl.insert(0, tblpr)
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'bottom'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '12')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), 'auto')
        borders.append(el)
    tblpr.append(borders)

    # --- 末尾空行 ---
    doc.add_paragraph()

    doc.save(output_path)
    print(f'已生成：{output_path}')


if __name__ == '__main__':
    file_path = '程序设计基础-教学大纲-2025秋.docx'
    kv_list = extract_table_kv(file_path)
    if kv_list:
        generate_docx(kv_list[0], 'test.docx')
