"""Single-module Word document export functions.

Each function takes a ReportRecord and returns a BytesIO buffer containing
a standalone .docx file for that module.
"""

from io import BytesIO
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def _setup_matplotlib_chinese():
    """Configure matplotlib for Chinese font rendering. Call once."""
    import matplotlib
    matplotlib.use('Agg')
    import matplotlib.pyplot as plt
    from matplotlib.font_manager import FontProperties
    # Try to find a Chinese font
    candidates = [
        'C:/Windows/Fonts/simsun.ttc',
        'C:/Windows/Fonts/SIMSUN.TTC',
        'C:/Windows/Fonts/msyh.ttc',
        'C:/Windows/Fonts/MSYH.TTC',
        'C:/Windows/Fonts/simhei.ttf',
        '/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc',
        '/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc',
        '/System/Library/Fonts/Songti.ttc',
    ]
    for path in candidates:
        import os
        if os.path.exists(path):
            return FontProperties(fname=path, size=10)
    return None


_chinese_font = None


def _get_chinese_font():
    global _chinese_font
    if _chinese_font is None:
        _chinese_font = _setup_matplotlib_chinese()
    return _chinese_font


def _generate_chart_image(section):
    """Generate a bar chart image for a grade section. Returns BytesIO or None."""
    import matplotlib.pyplot as plt

    segments = section.get('segments', [])
    if not segments:
        return None

    labels = [s['label'] for s in segments]
    counts = [s['count'] for s in segments]
    col_name = section.get('col_name', '')

    font = _get_chinese_font()
    fig, ax = plt.subplots(figsize=(7, 4))
    bars = ax.bar(labels, counts, color='#409eff', edgecolor='white')
    ax.set_title(f'学生{col_name}成绩分布图', fontproperties=font, fontsize=14)
    ax.set_xlabel('分数段', fontproperties=font, fontsize=11)
    ax.set_ylabel('人数', fontproperties=font, fontsize=11)

    # Set tick labels with font
    if font:
        for label in ax.get_xticklabels():
            label.set_fontproperties(font)
        for label in ax.get_yticklabels():
            label.set_fontproperties(font)

    # Add count labels on top of bars
    for bar, count in zip(bars, counts):
        ax.text(bar.get_x() + bar.get_width() / 2, bar.get_height() + 0.5,
                str(count), ha='center', va='bottom', fontsize=10)

    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    plt.tight_layout()

    buf = BytesIO()
    fig.savefig(buf, format='png', dpi=150)
    buf.seek(0)
    plt.close(fig)
    return buf


def _setup_document(report):
    """Create a fresh Document with default Song Ti font and centered title."""
    doc = Document()
    default_style = doc.styles['Normal']
    default_style.font.name = '宋体'
    default_style.font.size = Pt(12)

    title = doc.add_heading(report.report_name, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.name = '宋体'
        run.font.size = Pt(16)

    return doc


def _run_font(run, font_name='宋体', font_size=Pt(12)):
    run.font.name = font_name
    run.font.size = font_size
    return run


def _set_cell(cell, text, bold=False):
    cell.text = ''
    run = cell.paragraphs[0].add_run(str(text))
    _run_font(run)
    if bold:
        run.bold = True


def _add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        _run_font(run)
    return h


def _add_paragraph(doc, text):
    p = doc.add_paragraph(str(text))
    for run in p.runs:
        _run_font(run)
    return p


def _to_buffer(doc):
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ── Module 1: Course Basic Info Table ───────────────────────────────────────

def export_module_1_docx(report):
    doc = _setup_document(report)
    rd = report.report_data
    info = rd.get('module_1_course_info', {})
    _add_heading(doc, '一、课程基本信息表')

    table = doc.add_table(rows=7, cols=4, style='Table Grid')
    rows_data = [
        ['课程名称：', info.get('course_name', ''), '修课人数：', str(info.get('student_count', ''))],
        ['课程编号：', info.get('course_code', ''), '课序号：', str(info.get('course_seq', ''))],
        ['授课班级：', info.get('teaching_class', ''), '学时数：', str(info.get('total_hours', ''))],
        ['选用教材：', info.get('textbook', ''), '学分：', str(info.get('credits', ''))],
        ['开课院系：', info.get('department', ''), '授课教师：', str(info.get('teacher', ''))],
    ]
    for r_idx, (lbl1, val1, lbl2, val2) in enumerate(rows_data):
        _set_cell(table.cell(r_idx, 0), lbl1, bold=True)
        _set_cell(table.cell(r_idx, 1), val1)
        _set_cell(table.cell(r_idx, 2), lbl2, bold=True)
        _set_cell(table.cell(r_idx, 3), val2)

    _set_cell(table.cell(5, 0), '课程性质：', bold=True)
    table.cell(5, 1).merge(table.cell(5, 3))
    _set_cell(table.cell(5, 1), info.get('course_nature', ''))

    _set_cell(table.cell(6, 0), '课程类型：', bold=True)
    table.cell(6, 1).merge(table.cell(6, 3))
    _set_cell(table.cell(6, 1), info.get('course_type', ''))

    return _to_buffer(doc)


# ── Module 2: Course Objectives ─────────────────────────────────────────────

def export_module_2_docx(report):
    doc = _setup_document(report)
    _add_heading(doc, '二、课程目标')
    _add_paragraph(doc, report.report_data.get('module_2_objectives', '未填写'))
    return _to_buffer(doc)


# ── Module 3: Evaluation Standards ──────────────────────────────────────────

def export_module_3_docx(report):
    doc = _setup_document(report)
    _add_heading(doc, '三、课程评价标准')

    eval_data = report.report_data.get('module_3_evaluation_standards', '未填写')
    if isinstance(eval_data, list):
        for block in eval_data:
            if block['type'] == 'paragraph':
                _add_paragraph(doc, block['text'])
            elif block['type'] == 'table':
                grid = block['grid']
                num_cols = block['num_cols']
                if not grid or num_cols == 0:
                    continue
                num_rows = len(grid)
                tbl = doc.add_table(rows=num_rows, cols=num_cols, style='Table Grid')
                for r in range(num_rows):
                    for c in range(num_cols):
                        cell_data = grid[r][c]
                        if cell_data is None:
                            continue
                        cell = tbl.cell(r, c)
                        _set_cell(cell, cell_data['text'], bold=(r == 0))
                        if cell_data['colspan'] > 1 or cell_data['rowspan'] > 1:
                            end_r = r + cell_data['rowspan'] - 1
                            end_c = c + cell_data['colspan'] - 1
                            cell.merge(tbl.cell(end_r, end_c))
    else:
        _add_paragraph(doc, eval_data)

    return _to_buffer(doc)


# ── Module 4: Evaluation Results ────────────────────────────────────────────

def _render_module_4_to_docx(doc, module_4_data):
    """Render Module 4 sections into a docx Document. Shared by standalone export
    and merge export in views.py."""
    sections = module_4_data.get('sections')

    if sections:
        # Top-level fallback alert
        if module_4_data.get('fallback'):
            p = doc.add_paragraph()
            _run_font(p.add_run('⚠ '), font_size=Pt(10))
            _run_font(p.add_run(
                '未从课程大纲中提取到考核评价标准分数段定义，已使用默认分数段（优秀/良好/中等/及格/不及格）',
            ), font_size=Pt(10))

        for section in sections:
            # Build heading with weight info
            title = section['col_name']
            if section.get('weight_pct'):
                title += f'（占总评{section["weight_pct"]}%）'
            _add_heading(doc, title, level=2)

            # Normalization note
            if section.get('is_weighted'):
                p = doc.add_paragraph()
                _run_font(p.add_run('注：该列成绩已由折算后分数归一化至百分制'), font_size=Pt(10))

            # Segment source note
            src = section.get('segment_source', '')
            if src == 'fallback':
                p = doc.add_paragraph()
                _run_font(p.add_run('注：该列使用默认分数段'), font_size=Pt(10))
            elif src == 'default':
                p = doc.add_paragraph()
                _run_font(p.add_run('注：该列使用系统默认分数段'), font_size=Pt(10))

            _add_paragraph(doc, section['description_line_1'])
            _add_paragraph(doc, section['description_line_2'])

            segments = section['segments']
            num_cols = len(segments) + 2
            table = doc.add_table(rows=3, cols=num_cols, style='Table Grid')

            # Row 0: 类别 | segment labels... | 平均分
            _set_cell(table.cell(0, 0), '类别', bold=True)
            for ci, seg in enumerate(segments):
                _set_cell(table.cell(0, ci + 1), seg['label'], bold=True)
            _set_cell(table.cell(0, num_cols - 1), '平均分', bold=True)

            # Row 1: 人数 | counts... | avg_score
            _set_cell(table.cell(1, 0), '人数')
            for ci, seg in enumerate(segments):
                _set_cell(table.cell(1, ci + 1), str(seg['count']))
            _set_cell(table.cell(1, num_cols - 1), str(section['avg_score']))

            # Row 2: 比例% | pcts... | 100
            _set_cell(table.cell(2, 0), '比例%')
            for ci, seg in enumerate(segments):
                _set_cell(table.cell(2, ci + 1), f"{seg['pct']}%")
            _set_cell(table.cell(2, num_cols - 1), '100')

            # Chart image (best-effort, skip on failure)
            try:
                chart_buf = _generate_chart_image(section)
                if chart_buf:
                    chart_paragraph = doc.add_paragraph()
                    chart_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    chart_run = chart_paragraph.add_run()
                    chart_run.add_picture(chart_buf, width=Inches(5.5))
                    caption = doc.add_paragraph(f'图：学生{section["col_name"]}成绩分布图')
                    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in caption.runs:
                        _run_font(run, font_size=Pt(10))
            except Exception:
                pass  # chart generation is best-effort

            # Stats summary
            s = section['stats']
            p = doc.add_paragraph()
            for label, val in [
                ('参考人数：', s['count']),
                ('缺考人数：', s['missing']),
                ('最高分：', s['max']),
                ('最低分：', s['min']),
                ('平均分：', s['avg']),
                ('中位数：', s['median']),
                ('标准差：', s['stdev']),
            ]:
                _run_font(p.add_run(f'{label}{val}　　'))
            _run_font(p.add_run(f'及格率：{s["pass_rate"]}%'))

            # AI summary
            ai_summary = section.get('ai_summary', '')
            if ai_summary:
                _add_paragraph(doc, '')
                _add_paragraph(doc, ai_summary)
        return

    # Legacy fallback: old-format grade_analysis
    grade_analysis = module_4_data.get('grade_analysis', {})
    if not grade_analysis:
        _add_paragraph(doc, '暂无成绩数据')
        return

    for col_name, stats in grade_analysis.items():
        _add_heading(doc, col_name, level=2)

        p = doc.add_paragraph()
        for label, val in [
            ('参考人数：', stats['count']),
            ('缺考人数：', stats['missing']),
            ('最高分：', stats['max']),
            ('最低分：', stats['min']),
            ('平均分：', stats['avg']),
            ('中位数：', stats['median']),
            ('标准差：', stats['stdev']),
        ]:
            _run_font(p.add_run(f'{label}{val}　　'))
        _run_font(p.add_run(f'及格率：{stats["pass_rate"]}%'))

        dp = doc.add_paragraph('成绩分布：')
        for run in dp.runs:
            _run_font(run)

        dist_table = doc.add_table(rows=1 + len(stats['distribution']), cols=3, style='Table Grid')
        _set_cell(dist_table.cell(0, 0), '类别', bold=True)
        _set_cell(dist_table.cell(0, 1), '人数', bold=True)
        _set_cell(dist_table.cell(0, 2), '占比', bold=True)

        dist_idx = 1
        total = stats['count']
        for key, dist in stats['distribution'].items():
            _set_cell(dist_table.cell(dist_idx, 0), dist['label'])
            _set_cell(dist_table.cell(dist_idx, 1), str(dist['count']))
            pct = round(dist['count'] / total * 100, 1) if total else 0
            _set_cell(dist_table.cell(dist_idx, 2), f'{pct}%')
            dist_idx += 1


def export_module_4_docx(report):
    doc = _setup_document(report)
    _add_heading(doc, '四、课程评价结果')
    _render_module_4_to_docx(doc, report.report_data.get('module_4_evaluation_results', {}))
    return _to_buffer(doc)


# ── Module 5: Objective Achievement Degree ───────────────────────────────────

def _split_by_comma(val):
    """Split a value by comma into trimmed parts; returns [val] if no comma present.
    Matches the frontend splitByComma logic in Module5Editor.vue."""
    s = str(val) if val else ''
    if ',' not in s:
        return [s]
    return [p.strip() for p in s.split(',') if p.strip()]


def _flatten_section_5_1_rows(section):
    """Flatten section 5.1 objectives into table rows with comma-splitting.

    Matches frontend flattened51Rows computed property:
    each item whose name or weight_pct contains commas is expanded into
    multiple rows, with 课程目标 and 课程目标达成情况 merged per objective group.
    """
    objectives = section.get('objectives', [])
    rows = []
    for obj in objectives:
        obj_name = obj.get('name', '')
        achievement_rate = obj.get('achievement_rate', '0%')
        obj_rows = []
        for item in (obj.get('items') or []):
            item_parts = _split_by_comma(item.get('item', ''))
            weight_parts = _split_by_comma(item.get('weight_pct', '0%'))
            max_len = max(len(item_parts), len(weight_parts))
            for k in range(max_len):
                obj_rows.append({
                    'objective': obj_name,
                    'item': item_parts[k] if k < len(item_parts) else (item_parts[-1] if item_parts else ''),
                    'target_score': item.get('target_score', 0),
                    'avg_score': item.get('avg_score', 0),
                    'weight_pct': weight_parts[k] if k < len(weight_parts) else (weight_parts[-1] if weight_parts else '0%'),
                    'achievement_rate': achievement_rate,
                })
        rows.extend(obj_rows)
    return rows


def _render_section_5_1_to_docx(doc, data):
    """Render section 5.1 (课程目标达成情况) summary table into the document.

    Table structure:
    | 课程目标 | 评价内容 | 目标分值 | 平均得分 | 权重系数 | 课程目标达成情况 |
    With merged cells for 课程目标 and 课程目标达成情况 per objective group.
    Comma-separated item/weight fields are expanded into individual rows.
    """
    section_5_1 = data.get('section_5_1')
    if not section_5_1:
        return

    objectives = section_5_1.get('objectives', [])
    if not objectives:
        return

    # Heading and intro
    _add_heading(doc, '5.1 课程目标达成情况', level=2)
    _add_paragraph(doc, '课程目标达成情况计算见下表。')

    # Centered title line with semester
    title = section_5_1.get('title', '课程目标达成情况')
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(title)
    _run_font(title_run, font_size=Pt(12))
    title_run.bold = True

    # Flatten rows with comma-splitting (matches frontend flattened51Rows)
    all_rows = _flatten_section_5_1_rows(section_5_1)
    if not all_rows:
        return

    # Build table
    table = doc.add_table(rows=1 + len(all_rows), cols=6, style='Table Grid')

    # Header row
    headers = ['课程目标', '评价内容', '目标分值', '平均得分', '权重系数', '课程目标达成情况']
    for ci, h in enumerate(headers):
        _set_cell(table.cell(0, ci), h, bold=True)

    # Data rows with merged cells per objective group
    row_idx = 1
    i = 0
    while i < len(all_rows):
        obj_name = all_rows[i]['objective']
        achievement_rate = all_rows[i]['achievement_rate']
        start_row = row_idx
        first_of_group = True
        # Consume all rows belonging to this objective
        while i < len(all_rows) and all_rows[i]['objective'] == obj_name:
            r = all_rows[i]
            # Only set merged-column text on the first row; subsequent rows' cells
            # will be removed by merge() and their content would otherwise duplicate.
            if first_of_group:
                _set_cell(table.cell(row_idx, 0), obj_name)
                _set_cell(table.cell(row_idx, 5), achievement_rate)
                first_of_group = False
            _set_cell(table.cell(row_idx, 1), r['item'])
            _set_cell(table.cell(row_idx, 2), str(r['target_score']))
            _set_cell(table.cell(row_idx, 3), str(r['avg_score']))
            _set_cell(table.cell(row_idx, 4), r['weight_pct'])
            row_idx += 1
            i += 1

        end_row = row_idx - 1
        # Merge 课程目标 and 课程目标达成情况 columns if objective spans multiple rows
        if end_row > start_row:
            table.cell(start_row, 0).merge(table.cell(end_row, 0))
            table.cell(start_row, 5).merge(table.cell(end_row, 5))


def _render_module_5_to_docx(doc, data):
    """Render Module 5 (课程目标达成度) structured data into a document."""
    if not data or not data.get('generated'):
        _add_paragraph(doc, '暂无课程目标达成度数据。')
        return

    objectives = data.get('objectives', [])

    # Section 5.1: 课程目标达成情况 summary table
    _render_section_5_1_to_docx(doc, data)

    # Intro for existing content
    _add_paragraph(doc, '课程目标达成情况计算见表6。')

    # (1) 统计分析 — Table 6 (achievement calc) + Table 7 (distribution)
    achievement_table = data.get('achievement_table', [])
    distribution_table = data.get('distribution_table')

    if achievement_table:
        _add_heading(doc, '(1) 统计分析', level=2)

        report_title = data.get('report_title', '')
        _add_paragraph(doc, f'表6   {report_title}')

        from collections import OrderedDict
        obj_groups = OrderedDict()
        for row in achievement_table:
            obj = row.get('objective', '')
            if obj not in obj_groups:
                obj_groups[obj] = []
            obj_groups[obj].append(row)

        for obj_name, items in obj_groups.items():
            _add_heading(doc, obj_name, level=3)

            table = doc.add_table(rows=1 + len(items), cols=5, style='Table Grid')
            headers = ['评价内容', '目标分值', '平均得分', '权重系数', '课程目标达成情况']
            for ci, h in enumerate(headers):
                _set_cell(table.cell(0, ci), h, bold=True)

            # The achievement_rate is the same for all items; merge it across rows.
            ach_rate = items[0].get('achievement_rate', '') if items else ''
            for ri, row in enumerate(items):
                _set_cell(table.cell(ri + 1, 0), row.get('item', ''))
                _set_cell(table.cell(ri + 1, 1), str(row.get('target_score', '')))
                _set_cell(table.cell(ri + 1, 2), str(row.get('avg_score', '')))
                _set_cell(table.cell(ri + 1, 3), row.get('weight_pct', ''))
                # Only set achievement_rate on the first row to avoid text duplication
                if ri == 0:
                    _set_cell(table.cell(ri + 1, 4), ach_rate)

            if len(items) > 1:
                table.cell(1, 4).merge(table.cell(len(items), 4))

    if distribution_table and distribution_table.get('objectives'):
        _add_paragraph(doc, '表7  各课程目标的达成情况和分布表')

        obj_names = distribution_table['objectives']
        dist_rows = distribution_table.get('rows', [])

        num_obj_cols = len(obj_names) * 2
        num_cols = 1 + num_obj_cols

        # Single distribution table matching the web page
        dist_table = doc.add_table(rows=1 + len(dist_rows), cols=num_cols, style='Table Grid')

        # Header row: 达成度分布 | 目标1（人数）| 目标1（比例）| 目标2（人数）| 目标2（比例）| ...
        _set_cell(dist_table.cell(0, 0), '达成度分布', bold=True)
        for oi in range(len(obj_names)):
            _set_cell(dist_table.cell(0, 1 + oi * 2), f'{obj_names[oi]}（人数）', bold=True)
            _set_cell(dist_table.cell(0, 1 + oi * 2 + 1), f'{obj_names[oi]}（比例）', bold=True)

        for di, dist_row in enumerate(dist_rows):
            row_idx = 1 + di
            _set_cell(dist_table.cell(row_idx, 0), dist_row['label'])
            for oi in range(len(obj_names)):
                counts = dist_row.get('counts', [0] * len(obj_names))
                pcts = dist_row.get('pcts', [0] * len(obj_names))
                _set_cell(dist_table.cell(row_idx, 1 + oi * 2),
                         str(counts[oi] if oi < len(counts) else 0))
                _set_cell(dist_table.cell(row_idx, 1 + oi * 2 + 1),
                         f"{pcts[oi] if oi < len(pcts) else 0}%")

    # (2) 图形分析
    charts_added = False
    chart_num = 1

    # Chart 1: 课程目标达成情况
    try:
        chart1_buf = _generate_obj_achievement_summary_chart(data)
        if chart1_buf:
            if not charts_added:
                _add_heading(doc, '(2) 图形分析', level=2)
                charts_added = True
            chart_para = doc.add_paragraph()
            chart_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            chart_run = chart_para.add_run()
            chart_run.add_picture(chart1_buf, width=Inches(5))
            caption = doc.add_paragraph(f'图{chart_num}  课程目标达成情况')
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in caption.runs:
                _run_font(run, font_size=Pt(10))
            chart_num += 1
    except Exception:
        pass

    # Distribution charts per objective
    if distribution_table and distribution_table.get('objectives'):
        obj_names = distribution_table['objectives']
        for obj_name in obj_names:
            try:
                chart_buf = _generate_achievement_chart(obj_name, data)
                if chart_buf:
                    if not charts_added:
                        _add_heading(doc, '(2) 图形分析', level=2)
                        charts_added = True
                    chart_para = doc.add_paragraph()
                    chart_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    chart_run = chart_para.add_run()
                    chart_run.add_picture(chart_buf, width=Inches(5))
                    caption = doc.add_paragraph(f'图{chart_num}  {obj_name}达成度分布')
                    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in caption.runs:
                        _run_font(run, font_size=Pt(10))
                    chart_num += 1
            except Exception:
                pass

    # (3) 教学班整体课程目标达成情况分析
    per_objective_analysis = data.get('per_objective_analysis', [])
    if per_objective_analysis:
        _add_heading(doc, '(3) 教学班整体课程目标达成情况分析', level=2)
        labels = ['①', '②', '③', '④', '⑤', '⑥', '⑦', '⑧']
        for i, poa in enumerate(per_objective_analysis):
            label = labels[i] if i < len(labels) else f'({i+1})'
            _add_heading(doc, f'{label}{poa.get("objective", "")}的达成情况分析', level=3)
            analysis_text = poa.get('analysis', '')
            if analysis_text:
                _add_paragraph(doc, analysis_text)

    # (4) 学生个体课程目标达成情况分析
    low_students = data.get('low_students', [])
    if low_students:
        _add_heading(doc, '(4) 学生个体课程目标达成情况分析', level=2)
        _add_heading(doc, '需要关注及跟踪的学生', level=3)
        _add_paragraph(doc, '表8  未达成课程目标学生汇总表')

        num_obj = len(objectives) if objectives else 1
        table = doc.add_table(rows=1 + len(low_students), cols=1 + num_obj, style='Table Grid')
        _set_cell(table.cell(0, 0), '姓名', bold=True)
        for oi, obj_name in enumerate(objectives):
            _set_cell(table.cell(0, 1 + oi), f'{obj_name}达成度', bold=True)

        for ri, student in enumerate(low_students):
            _set_cell(table.cell(ri + 1, 0), student['name'])
            achievements = student.get('achievements', {})
            for oi, obj_name in enumerate(objectives):
                val = achievements.get(obj_name, 0)
                _set_cell(table.cell(ri + 1, 1 + oi), str(val))

        _add_paragraph(doc, '以上几位学生在课程目标的达成度上较低，需要重点关注与跟踪。')

    # Student scatter chart (after section 4)
    try:
        scatter_buf = _generate_student_scatter_chart(data)
        if scatter_buf:
            chart_para = doc.add_paragraph()
            chart_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            chart_run = chart_para.add_run()
            chart_run.add_picture(scatter_buf, width=Inches(5.5))
            caption = doc.add_paragraph(f'图{chart_num}  学生个体课程目标达成度分布')
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in caption.runs:
                _run_font(run, font_size=Pt(10))
    except Exception:
        pass


def _generate_achievement_chart(obj_name, data):
    """Generate a bar chart for a single objective's achievement distribution.

    Shows 比例 (%) on the Y-axis, matching the web page format.
    """
    import matplotlib.pyplot as plt

    distribution = data.get('distribution_table')
    if not distribution:
        return None

    obj_names = distribution.get('objectives', [])
    if obj_name not in obj_names:
        return None

    oi = obj_names.index(obj_name)
    dists = distribution.get('rows', [])
    if not dists:
        return None

    labels = [d['label'] for d in dists]
    pcts = [d.get('pcts', [0]*len(obj_names))[oi] for d in dists]

    obj_colors = ['#67c23a', '#e6a23c', '#f56c6c', '#909399']
    bar_color = obj_colors[oi % len(obj_colors)]

    font = _get_chinese_font()
    fig, ax = plt.subplots(figsize=(6, 4))
    bars = ax.bar(labels, pcts, color=bar_color, edgecolor='white')
    ax.set_title(f'{obj_name}达成情况分布', fontproperties=font, fontsize=13)
    ax.set_xlabel('达成度区间', fontproperties=font, fontsize=10)
    ax.set_ylabel('比例 (%)', fontproperties=font, fontsize=10)

    if font:
        for label in ax.get_xticklabels():
            label.set_fontproperties(font)
        for label in ax.get_yticklabels():
            label.set_fontproperties(font)

    for bar, pct in zip(bars, pcts):
        ax.text(bar.get_x() + bar.get_width() / 2, bar.get_height() + 0.5,
                f'{pct}%', ha='center', va='bottom', fontsize=9)

    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    plt.tight_layout()

    buf = BytesIO()
    fig.savefig(buf, format='png', dpi=150)
    buf.seek(0)
    plt.close(fig)
    return buf


def _generate_obj_achievement_summary_chart(data):
    """Generate Chart 1: 课程目标达成情况 — bar chart of objective achievement rates."""
    import matplotlib.pyplot as plt

    section = data.get('section_5_1')
    if not section:
        return None
    objectives = section.get('objectives', [])
    if not objectives:
        return None

    names = [o.get('name', '') for o in objectives]
    rates_str = [o.get('achievement_rate', '0%') for o in objectives]
    rates = []
    for r in rates_str:
        try:
            rates.append(float(str(r).replace('%', '').strip()))
        except (ValueError, TypeError):
            rates.append(0.0)

    font = _get_chinese_font()
    fig, ax = plt.subplots(figsize=(6, 4))
    bars = ax.bar(names, rates, color='#409eff', edgecolor='white', width=0.5)
    ax.set_title('课程目标达成情况', fontproperties=font, fontsize=13)
    ax.set_ylabel('课程目标达成情况 (%)', fontproperties=font, fontsize=10)

    if font:
        for label in ax.get_xticklabels():
            label.set_fontproperties(font)
        for label in ax.get_yticklabels():
            label.set_fontproperties(font)

    ax.set_ylim(0, 110)
    for bar, rate in zip(bars, rates):
        ax.text(bar.get_x() + bar.get_width() / 2, bar.get_height() + 1,
                f'{rate}%', ha='center', va='bottom', fontsize=10)

    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    plt.tight_layout()

    buf = BytesIO()
    fig.savefig(buf, format='png', dpi=150)
    buf.seek(0)
    plt.close(fig)
    return buf


def _generate_student_scatter_chart(data):
    """Generate scatter chart: 学生个体课程目标达成度分布 with average line."""
    import matplotlib.pyplot as plt

    students = data.get('student_achievements', [])
    if not students:
        return None
    overall_avg = data.get('overall_avg_achievement', 0)

    x = list(range(1, len(students) + 1))
    y = [s.get('avg_achievement', 0) for s in students]

    font = _get_chinese_font()
    fig, ax = plt.subplots(figsize=(8, 4))
    ax.scatter(x, y, s=20, c='#409eff', alpha=0.6, edgecolors='none')
    ax.axhline(y=overall_avg, color='red', linestyle='-', linewidth=1.5,
               label=f'总平均达成度: {overall_avg}%')

    ax.set_title('学生个体课程目标达成度分布', fontproperties=font, fontsize=13)
    ax.set_xlabel('学生序号', fontproperties=font, fontsize=10)
    ax.set_ylabel('平均达成度 (%)', fontproperties=font, fontsize=10)

    if font:
        for label in ax.get_xticklabels():
            label.set_fontproperties(font)
        for label in ax.get_yticklabels():
            label.set_fontproperties(font)

    ax.legend(prop=font if font else None)
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    plt.tight_layout()

    buf = BytesIO()
    fig.savefig(buf, format='png', dpi=150)
    buf.seek(0)
    plt.close(fig)
    return buf


def export_module_5_docx(report):
    doc = _setup_document(report)
    _add_heading(doc, '五、课程目标达成度')
    _render_module_5_to_docx(doc, report.report_data.get('module_5_objective_achievement', {}))
    return _to_buffer(doc)


# ── Module 6: Improvement Plan ──────────────────────────────────────────────

def export_module_6_docx(report):
    doc = _setup_document(report)
    _add_heading(doc, '六、课程持续改进方案及措施')

    data = report.report_data.get('module_6_improvement_plan', {})

    if isinstance(data, dict):
        # New structured format
        # 4.1 (now 6.1)
        _add_heading(doc, '6.1 连续两年课程评价结果系统地纳入课程持续改进的措施及其效果描述', level=2)
        _add_paragraph(doc, data.get('part1', ''))

        # 4.2 (now 6.2)
        _add_heading(doc, '6.2 本年度课程教学环节发现的问题、相应持续改进的措施以及描述预期将可能达到的效果', level=2)
        part2 = data.get('part2', {})
        if isinstance(part2, dict):
            _add_heading(doc, '(1) 存在的问题', level=3)
            _add_paragraph(doc, part2.get('problems', ''))
            _add_heading(doc, '(2) 持续改进措施', level=3)
            _add_paragraph(doc, part2.get('measures', ''))
            _add_heading(doc, '(3) 预期效果', level=3)
            _add_paragraph(doc, part2.get('expected_effects', ''))
        else:
            _add_paragraph(doc, str(part2))

        # 4.3 (now 6.3)
        _add_heading(doc, '6.3 其他可用的协助持续改进的资源', level=2)
        _add_paragraph(doc, data.get('part3', ''))
    else:
        # Legacy string format
        _add_paragraph(doc, str(data) if data else '待后续版本实现')

    return _to_buffer(doc)


def _render_module_6_to_docx(doc, data):
    """Render Module 6 structured data into an existing document (for merge)."""
    if isinstance(data, dict):
        _add_heading(doc, '6.1 连续两年课程评价结果系统地纳入课程持续改进的措施及其效果描述', level=2)
        _add_paragraph(doc, data.get('part1', ''))

        _add_heading(doc, '6.2 本年度课程教学环节发现的问题、相应持续改进的措施以及描述预期将可能达到的效果', level=2)
        part2 = data.get('part2', {})
        if isinstance(part2, dict):
            _add_heading(doc, '(1) 存在的问题', level=3)
            _add_paragraph(doc, part2.get('problems', ''))
            _add_heading(doc, '(2) 持续改进措施', level=3)
            _add_paragraph(doc, part2.get('measures', ''))
            _add_heading(doc, '(3) 预期效果', level=3)
            _add_paragraph(doc, part2.get('expected_effects', ''))
        else:
            _add_paragraph(doc, str(part2))

        _add_heading(doc, '6.3 其他可用的协助持续改进的资源', level=2)
        _add_paragraph(doc, data.get('part3', ''))
    else:
        _add_paragraph(doc, str(data) if data else '待后续版本实现')
