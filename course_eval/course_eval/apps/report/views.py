from rest_framework.views import APIView
from rest_framework.parsers import MultiPartParser

from course_eval.utils.response import api_response
from .models import ReportRecord
from .serializers import ReportRecordSerializer, ReportGenerateSerializer, ModuleUpdateSerializer
from .report_engine import (
    generate_report, generate_module_1, generate_module_2, generate_module_3,
    generate_module_4, generate_module_5, generate_module_6, _parse_syllabus, _analyze_student_info,
)
from .report_engine import _find_file
from .module_docx import (
    export_module_1_docx, export_module_2_docx, export_module_3_docx,
    export_module_4_docx, export_module_5_docx, export_module_6_docx,
    _render_module_4_to_docx, _render_module_5_to_docx, _render_module_6_to_docx,
)
from .excel_engine import generate_achievement_excel
from .pdf_engine import generate_pdf

MODULE_MAP = {
    1: ('module_1_course_info', 'module_1_status'),
    2: ('module_2_objectives', 'module_2_status'),
    3: ('module_3_evaluation_standards', 'module_3_status'),
    4: ('module_4_evaluation_results', 'module_4_status'),
    5: ('module_5_objective_achievement', 'module_5_status'),
    6: ('module_6_improvement_plan', 'module_6_status'),
}

EXPORT_FUNCTIONS = {
    1: export_module_1_docx,
    2: export_module_2_docx,
    3: export_module_3_docx,
    4: export_module_4_docx,
    5: export_module_5_docx,
    6: export_module_6_docx,
}


def _get_report_or_404(pk, user):
    report = ReportRecord.objects.filter(id=pk, user=user).first()
    if not report:
        return None
    return report


def _regenerate_module(report, module_num):
    """Re-parse source files and regenerate a single module's data."""
    user_id = report.user_id
    course = report.course
    class_group = report.class_group
    semester = report.semester

    if module_num == 1:
        syllabus_fields = _parse_syllabus(report.syllabus_file)
        raw_table_kv = syllabus_fields.pop('_all_table_kv', {})
        student_stats = {}
        if report.student_info_file:
            from course_eval.apps.excel.data_handler import get_effective_data
            info_data = get_effective_data(report.student_info_file, user_id)
            student_stats = _analyze_student_info(info_data['headers'], info_data['rows'])
        return generate_module_1(course, class_group, syllabus_fields, student_stats, raw_table_kv)

    elif module_num == 2:
        syllabus_fields = _parse_syllabus(report.syllabus_file)
        return generate_module_2(syllabus_fields)

    elif module_num == 3:
        syllabus_fields = _parse_syllabus(report.syllabus_file)
        return generate_module_3(syllabus_fields)

    elif module_num == 4:
        syllabus_fields = _parse_syllabus(report.syllabus_file)
        evaluation_standards = syllabus_fields.get('evaluation_standards')
        return generate_module_4(report.grades_file, user_id, evaluation_standards)

    elif module_num == 5:
        from .report_engine import generate_module_5
        syllabus_fields = _parse_syllabus(report.syllabus_file)
        evaluation_standards = syllabus_fields.get('evaluation_standards')
        objectives = report.report_data.get('module_2_objectives', '')
        return generate_module_5(
            course.name, objectives, evaluation_standards,
            report.grades_file, user_id,
            semester_name=semester.name,
        )

    elif module_num == 6:
        from .report_engine import _build_module6_context
        rd = report.report_data
        module_2 = rd.get('module_2_objectives', '')
        module_4 = rd.get('module_4_evaluation_results', {})
        context = _build_module6_context(course.name, module_2, module_4)
        return generate_module_6(context)

    return None


# ── Existing Views ──────────────────────────────────────────────────────────

class ReportGenerateView(APIView):
    def post(self, request):
        serializer = ReportGenerateSerializer(data=request.data)
        if not serializer.is_valid():
            return api_response(code=400, msg='参数错误', data=serializer.errors, http_status=400)

        data = serializer.validated_data
        user_id = request.user.id

        result = generate_report(
            user_id, data['course_id'], data['class_id'], data['semester_name']
        )

        # Upsert — one report per course/class/semester/user
        report, _ = ReportRecord.objects.update_or_create(
            course=result['course'],
            class_group=result['class_group'],
            semester=result['semester'],
            user=request.user,
            defaults={
                'report_name': result['report_name'],
                'report_data': result['report_data'],
                'syllabus_file': result['syllabus_file'],
                'student_info_file': result['student_info_file'],
                'grades_file': result['grades_file'],
                'status': 'completed',
                'module_1_status': 'draft',
                'module_2_status': 'draft',
                'module_3_status': 'draft',
                'module_4_status': 'draft',
                'module_5_status': 'draft',
                'module_6_status': 'draft',
            },
        )

        return api_response(data=ReportRecordSerializer(report).data)


class ReportPreviewView(APIView):
    def get(self, request, pk):
        report = _get_report_or_404(pk, request.user)
        if not report:
            return api_response(code=404, msg='报告不存在', http_status=404)

        return api_response(data={
            **ReportRecordSerializer(report).data,
            'report_data': report.report_data,
        })


class ReportExportView(APIView):
    def get(self, request, pk):
        report = _get_report_or_404(pk, request.user)
        if not report:
            return api_response(code=404, msg='报告不存在', http_status=404)

        export_format = request.query_params.get('export_format', 'docx')

        if export_format == 'pdf':
            return self._export_pdf(report)
        return self._export_docx(report)

    def _export_docx(self, report):
        from django.http import FileResponse
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        rd = report.report_data
        doc = Document()

        default_style = doc.styles['Normal']
        default_style.font.name = '宋体'
        default_style.font.size = Pt(12)

        def _run_font(run, font_name='宋体', font_size=Pt(12)):
            run.font.name = font_name
            run.font.size = font_size
            return run

        def _set_cell(cell, text, bold=False):
            cell.text = ''
            run = cell.paragraphs[0].add_run(str(text))
            _run_font(run)
            if bold:
                run.bold = True

        title = doc.add_heading(report.report_name, level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in title.runs:
            _run_font(run)

        # Module 1
        h = doc.add_heading('一、课程基本信息表', level=1)
        for run in h.runs:
            _run_font(run)
        info = rd.get('module_1_course_info', {})
        table = doc.add_table(rows=7, cols=4, style='Table Grid')
        rows_data = [
            ['课程名称：', info.get('course_name', ''), '修课人数：', str(info.get('student_count', ''))],
            ['课程编号：', info.get('course_code', ''), '课序号：', str(info.get('course_seq', ''))],
            ['授课班级：', info.get('teaching_class', ''), '学时数：', str(info.get('total_hours', ''))],
            ['选用教材：', info.get('textbook', ''), '学分：', str(info.get('credits', ''))],
            ['开课院系：', info.get('department', ''), '授课教师：', str(info.get('teacher', ''))],
        ]
        for r_idx, (lbl1, val1, lbl2, val2) in enumerate(rows_data):
            _set_cell(table.cell(r_idx, 0), lbl1, bold=True)
            _set_cell(table.cell(r_idx, 1), val1)
            _set_cell(table.cell(r_idx, 2), lbl2, bold=True)
            _set_cell(table.cell(r_idx, 3), val2)

        _set_cell(table.cell(5, 0), '课程性质：', bold=True)
        table.cell(5, 1).merge(table.cell(5, 3))
        _set_cell(table.cell(5, 1), info.get('course_nature', ''))

        _set_cell(table.cell(6, 0), '课程类型：', bold=True)
        table.cell(6, 1).merge(table.cell(6, 3))
        _set_cell(table.cell(6, 1), info.get('course_type', ''))

        # Module 2
        h = doc.add_heading('二、课程目标', level=1)
        for run in h.runs:
            _run_font(run)
        p = doc.add_paragraph(rd.get('module_2_objectives', '未填写'))
        for run in p.runs:
            _run_font(run)

        # Module 3
        h = doc.add_heading('三、课程评价标准', level=1)
        for run in h.runs:
            _run_font(run)
        eval_data = rd.get('module_3_evaluation_standards', '未填写')
        if isinstance(eval_data, list):
            for block in eval_data:
                if block['type'] == 'paragraph':
                    p = doc.add_paragraph(block['text'])
                    for run in p.runs:
                        _run_font(run)
                elif block['type'] == 'table':
                    grid = block['grid']
                    num_cols = block['num_cols']
                    if not grid or num_cols == 0:
                        continue
                    num_rows = len(grid)
                    tbl = doc.add_table(rows=num_rows, cols=num_cols, style='Table Grid')
                    for r in range(num_rows):
                        for c in range(num_cols):
                            cell_data = grid[r][c]
                            if cell_data is None:
                                continue
                            cell = tbl.cell(r, c)
                            _set_cell(cell, cell_data['text'], bold=(r == 0))
                            if cell_data['colspan'] > 1 or cell_data['rowspan'] > 1:
                                end_r = r + cell_data['rowspan'] - 1
                                end_c = c + cell_data['colspan'] - 1
                                cell.merge(tbl.cell(end_r, end_c))
        else:
            p = doc.add_paragraph(eval_data)
            for run in p.runs:
                _run_font(run)

        # Module 4
        h = doc.add_heading('四、课程评价结果', level=1)
        for run in h.runs:
            _run_font(run)
        eval_results = rd.get('module_4_evaluation_results', {})
        grade_analysis = eval_results.get('grade_analysis', {})
        if grade_analysis:
            for col_name, stats in grade_analysis.items():
                h2 = doc.add_heading(col_name, level=2)
                for run in h2.runs:
                    _run_font(run)
                p = doc.add_paragraph()
                for label, val in [
                    ('参考人数：', stats['count']),
                    ('缺考人数：', stats['missing']),
                    ('最高分：', stats['max']),
                    ('最低分：', stats['min']),
                    ('平均分：', stats['avg']),
                    ('中位数：', stats['median']),
                    ('标准差：', stats['stdev']),
                ]:
                    _run_font(p.add_run(f'{label}{val}　　'))
                _run_font(p.add_run(f'及格率：{stats["pass_rate"]}%'))

                dp = doc.add_paragraph('成绩分布：')
                for run in dp.runs:
                    _run_font(run)
                dist_table = doc.add_table(rows=1 + len(stats['distribution']), cols=3, style='Table Grid')
                _set_cell(dist_table.cell(0, 0), '类别', bold=True)
                _set_cell(dist_table.cell(0, 1), '人数', bold=True)
                _set_cell(dist_table.cell(0, 2), '占比', bold=True)
                dist_idx = 1
                total = stats['count']
                for key, dist in stats['distribution'].items():
                    _set_cell(dist_table.cell(dist_idx, 0), dist['label'])
                    _set_cell(dist_table.cell(dist_idx, 1), str(dist['count']))
                    pct = round(dist['count'] / total * 100, 1) if total else 0
                    _set_cell(dist_table.cell(dist_idx, 2), f'{pct}%')
                    dist_idx += 1
        else:
            p = doc.add_paragraph('暂无成绩数据')
            for run in p.runs:
                _run_font(run)

        # Module 5: 课程目标达成度
        h = doc.add_heading('五、课程目标达成度', level=1)
        for run in h.runs:
            _run_font(run)
        _render_module_5_to_docx(doc, rd.get('module_5_objective_achievement', {}))

        # Module 6: 持续改进方案
        h = doc.add_heading('六、课程持续改进方案及措施', level=1)
        for run in h.runs:
            _run_font(run)
        _render_module_6_to_docx(doc, rd.get('module_6_improvement_plan', {}))

        # Signature lines
        from datetime import date
        today = date.today()
        today_str = f'{today.year}.{today.month}.{today.day}'
        for label in ('任课教师签名：', '责任教授或系主任签名：'):
            sig_p = doc.add_paragraph()
            _run_font(sig_p.add_run(f'{label}                '))
            _run_font(sig_p.add_run(f'日期：{today_str}'))

        from io import BytesIO

        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)

        report.status = 'exported'
        report.save(update_fields=['status'])

        return FileResponse(
            buf,
            as_attachment=True,
            filename=f'{report.report_name}.docx',
        )

    def _export_pdf(self, report):
        from django.http import FileResponse

        buf = generate_pdf(report.report_name, report.report_data)

        report.status = 'exported'
        report.save(update_fields=['status'])

        return FileResponse(
            buf,
            as_attachment=True,
            filename=f'{report.report_name}.pdf',
            content_type='application/pdf',
        )


class ReportListView(APIView):
    def get(self, request):
        course_id = request.query_params.get('course_id')
        class_id = request.query_params.get('class_id')
        semester_name = request.query_params.get('semester_name')

        qs = ReportRecord.objects.filter(user=request.user)
        if course_id:
            qs = qs.filter(course_id=course_id)
        if class_id:
            qs = qs.filter(class_group_id=class_id)
        if semester_name:
            qs = qs.filter(semester__name=semester_name)

        serializer = ReportRecordSerializer(qs, many=True)
        return api_response(data=serializer.data)


# ── New per-module views ────────────────────────────────────────────────────

class ModuleGenerateView(APIView):
    """Generate (or re-generate) a single module's data."""

    def post(self, request, pk, module_num):
        if module_num not in MODULE_MAP:
            return api_response(code=400, msg='无效的模块编号', http_status=400)

        report = _get_report_or_404(pk, request.user)
        if not report:
            return api_response(code=404, msg='报告不存在', http_status=404)

        try:
            new_data = _regenerate_module(report, module_num)
        except Exception as e:
            return api_response(code=500, msg=f'模块生成失败: {str(e)}', http_status=500)

        # Preserve Excel generation flags when regenerating Module 5
        if module_num == 5:
            old_data = report.report_data.get(MODULE_MAP[5][0], {}) or {}
            for flag in ('excel_generated', 'excel_filename'):
                if flag in old_data:
                    new_data[flag] = old_data[flag]

        report.report_data[MODULE_MAP[module_num][0]] = new_data
        setattr(report, MODULE_MAP[module_num][1], 'draft')
        report.save(update_fields=['report_data', MODULE_MAP[module_num][1]])

        return api_response(data={
            'module_data': new_data,
            'module_status': 'draft',
        })


class ModuleUpdateView(APIView):
    """Save user-edited module data, optionally confirming it."""

    def put(self, request, pk, module_num):
        if module_num not in MODULE_MAP:
            return api_response(code=400, msg='无效的模块编号', http_status=400)

        report = _get_report_or_404(pk, request.user)
        if not report:
            return api_response(code=404, msg='报告不存在', http_status=404)

        serializer = ModuleUpdateSerializer(data=request.data)
        if not serializer.is_valid():
            return api_response(code=400, msg='参数错误', data=serializer.errors, http_status=400)

        validated = serializer.validated_data
        module_key = MODULE_MAP[module_num][0]
        status_field = MODULE_MAP[module_num][1]
        new_status = 'confirmed' if validated['confirmed'] else 'draft'

        report.report_data[module_key] = validated['data']
        setattr(report, status_field, new_status)
        report.save(update_fields=['report_data', status_field])

        return api_response(data={
            'module_data': validated['data'],
            'module_status': new_status,
        })


class ModuleExportView(APIView):
    """Export a single module as a standalone Word document."""

    def get(self, request, pk, module_num):
        if module_num not in EXPORT_FUNCTIONS:
            return api_response(code=400, msg='无效的模块编号', http_status=400)

        report = _get_report_or_404(pk, request.user)
        if not report:
            return api_response(code=404, msg='报告不存在', http_status=404)

        try:
            buf = EXPORT_FUNCTIONS[module_num](report)
        except Exception as e:
            return api_response(code=500, msg=f'模块导出失败: {str(e)}', http_status=500)

        from django.http import FileResponse
        module_data_key = MODULE_MAP[module_num][0]
        module_label = ['', '一', '二', '三', '四', '五', '六'][module_num]
        filename = f'{report.report_name}_{module_label}_{module_data_key}.docx'

        return FileResponse(
            buf,
            as_attachment=True,
            filename=filename,
        )


class ReportMergeView(APIView):
    """Merge all confirmed modules into a final combined report."""

    def get(self, request, pk):
        report = _get_report_or_404(pk, request.user)
        if not report:
            return api_response(code=404, msg='报告不存在', http_status=404)

        # Check all modules are confirmed
        unconfirmed = []
        for num, (_, status_field) in MODULE_MAP.items():
            if getattr(report, status_field) != 'confirmed':
                unconfirmed.append(f'模块{num}')

        if unconfirmed:
            return api_response(
                code=400,
                msg=f'以下模块尚未确认: {", ".join(unconfirmed)}',
                data={'unconfirmed_modules': [int(m.replace("模块", "")) for m in unconfirmed]},
                http_status=400,
            )

        export_format = request.query_params.get('export_format', 'docx')

        if export_format == 'pdf':
            return self._merge_pdf(report)
        return self._merge_docx(report)

    def _merge_docx(self, report):
        from django.http import FileResponse
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        rd = report.report_data
        doc = Document()

        default_style = doc.styles['Normal']
        default_style.font.name = '宋体'
        default_style.font.size = Pt(12)

        def _run_font(run, font_name='宋体', font_size=Pt(12)):
            run.font.name = font_name
            run.font.size = font_size
            return run

        def _set_cell(cell, text, bold=False):
            cell.text = ''
            run = cell.paragraphs[0].add_run(str(text))
            _run_font(run)
            if bold:
                run.bold = True

        title = doc.add_heading(report.report_name, level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in title.runs:
            _run_font(run)

        # ── Module 1 ──
        h = doc.add_heading('一、课程基本信息表', level=1)
        for run in h.runs:
            _run_font(run)
        info = rd.get('module_1_course_info', {})
        table = doc.add_table(rows=7, cols=4, style='Table Grid')
        rows_data = [
            ['课程名称：', info.get('course_name', ''), '修课人数：', str(info.get('student_count', ''))],
            ['课程编号：', info.get('course_code', ''), '课序号：', str(info.get('course_seq', ''))],
            ['授课班级：', info.get('teaching_class', ''), '学时数：', str(info.get('total_hours', ''))],
            ['选用教材：', info.get('textbook', ''), '学分：', str(info.get('credits', ''))],
            ['开课院系：', info.get('department', ''), '授课教师：', str(info.get('teacher', ''))],
        ]
        for r_idx, (lbl1, val1, lbl2, val2) in enumerate(rows_data):
            _set_cell(table.cell(r_idx, 0), lbl1, bold=True)
            _set_cell(table.cell(r_idx, 1), val1)
            _set_cell(table.cell(r_idx, 2), lbl2, bold=True)
            _set_cell(table.cell(r_idx, 3), val2)

        _set_cell(table.cell(5, 0), '课程性质：', bold=True)
        table.cell(5, 1).merge(table.cell(5, 3))
        _set_cell(table.cell(5, 1), info.get('course_nature', ''))

        _set_cell(table.cell(6, 0), '课程类型：', bold=True)
        table.cell(6, 1).merge(table.cell(6, 3))
        _set_cell(table.cell(6, 1), info.get('course_type', ''))

        # ── Module 2 ──
        h = doc.add_heading('二、课程目标', level=1)
        for run in h.runs:
            _run_font(run)
        p = doc.add_paragraph(rd.get('module_2_objectives', '未填写'))
        for run in p.runs:
            _run_font(run)

        # ── Module 3 ──
        h = doc.add_heading('三、课程评价标准', level=1)
        for run in h.runs:
            _run_font(run)
        eval_data = rd.get('module_3_evaluation_standards', '未填写')
        if isinstance(eval_data, list):
            for block in eval_data:
                if block['type'] == 'paragraph':
                    p = doc.add_paragraph(block['text'])
                    for run in p.runs:
                        _run_font(run)
                elif block['type'] == 'table':
                    grid = block['grid']
                    num_cols = block['num_cols']
                    if not grid or num_cols == 0:
                        continue
                    num_rows = len(grid)
                    tbl = doc.add_table(rows=num_rows, cols=num_cols, style='Table Grid')
                    for r in range(num_rows):
                        for c in range(num_cols):
                            cell_data = grid[r][c]
                            if cell_data is None:
                                continue
                            cell = tbl.cell(r, c)
                            _set_cell(cell, cell_data['text'], bold=(r == 0))
                            if cell_data['colspan'] > 1 or cell_data['rowspan'] > 1:
                                end_r = r + cell_data['rowspan'] - 1
                                end_c = c + cell_data['colspan'] - 1
                                cell.merge(tbl.cell(end_r, end_c))
        else:
            p = doc.add_paragraph(eval_data)
            for run in p.runs:
                _run_font(run)

        # ── Module 4 ──
        h = doc.add_heading('四、课程评价结果', level=1)
        for run in h.runs:
            _run_font(run)
        _render_module_4_to_docx(doc, rd.get('module_4_evaluation_results', {}))

        # ── Module 5: 课程目标达成度 ──
        h = doc.add_heading('五、课程目标达成度', level=1)
        for run in h.runs:
            _run_font(run)
        _render_module_5_to_docx(doc, rd.get('module_5_objective_achievement', {}))

        # ── Module 6: 持续改进方案 ──
        h = doc.add_heading('六、课程持续改进方案及措施', level=1)
        for run in h.runs:
            _run_font(run)
        _render_module_6_to_docx(doc, rd.get('module_6_improvement_plan', {}))

        # Signature lines
        from datetime import date
        today = date.today()
        today_str = f'{today.year}.{today.month}.{today.day}'
        for label in ('任课教师签名：', '责任教授或系主任签名：'):
            sig_p = doc.add_paragraph()
            _run_font(sig_p.add_run(f'{label}                '))
            _run_font(sig_p.add_run(f'日期：{today_str}'))

        from io import BytesIO

        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)

        report.status = 'exported'
        report.save(update_fields=['status'])

        return FileResponse(
            buf,
            as_attachment=True,
            filename=f'{report.report_name}.docx',
        )

    def _merge_pdf(self, report):
        from django.http import FileResponse

        buf = generate_pdf(report.report_name, report.report_data)

        report.status = 'exported'
        report.save(update_fields=['status'])

        return FileResponse(
            buf,
            as_attachment=True,
            filename=f'{report.report_name}.pdf',
            content_type='application/pdf',
        )


class Module5ExcelGenerateView(APIView):
    """Generate the course objective achievement Excel for Module 5."""

    def post(self, request, pk):
        report = _get_report_or_404(pk, request.user)
        if not report:
            return api_response(code=404, msg='报告不存在', http_status=404)

        try:
            buf = generate_achievement_excel(report)
        except Exception as e:
            return api_response(code=500, msg=f'Excel生成失败: {str(e)}', http_status=500)

        if buf is None:
            return api_response(code=400, msg='成绩数据或教学大纲不足，无法生成Excel', http_status=400)

        # Save file to media/excel/
        import os
        from django.conf import settings
        excel_dir = os.path.join(settings.MEDIA_ROOT, 'excel')
        os.makedirs(excel_dir, exist_ok=True)

        sem_name = report.semester.name if report.semester else ''
        course_name = report.course.name if report.course else ''
        filename = f'{sem_name}{course_name}课程达成度计算.xlsx'
        filepath = os.path.join(excel_dir, filename)

        with open(filepath, 'wb') as f:
            f.write(buf.getvalue())

        # Store flags in report_data
        module_data = report.report_data.get('module_5_objective_achievement', {}) or {}
        module_data['excel_generated'] = True
        module_data['excel_filename'] = filename
        report.report_data['module_5_objective_achievement'] = module_data
        report.save(update_fields=['report_data'])

        return api_response(data={
            'success': True,
            'filename': filename,
        })


class Module5ExcelDownloadView(APIView):
    """Download the generated achievement Excel for Module 5."""

    def get(self, request, pk):
        report = _get_report_or_404(pk, request.user)
        if not report:
            return api_response(code=404, msg='报告不存在', http_status=404)

        module_data = report.report_data.get('module_5_objective_achievement', {}) or {}
        filename = module_data.get('excel_filename', '')
        if not filename:
            return api_response(code=404, msg='Excel文件尚未生成', http_status=404)

        import os
        from django.conf import settings
        from django.http import FileResponse

        filepath = os.path.join(settings.MEDIA_ROOT, 'excel', filename)
        if not os.path.exists(filepath):
            return api_response(code=404, msg='Excel文件不存在，请重新生成', http_status=404)

        return FileResponse(
            open(filepath, 'rb'),
            as_attachment=True,
            filename=filename,
            content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        )
