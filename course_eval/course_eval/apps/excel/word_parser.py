import os
import re
from difflib import SequenceMatcher
from docx import Document
from docx.oxml.shared import qn

# ── Wingdings checkbox symbol mapping ──────────────────────────────────────
_SYM_CHECKED = '☑'
_SYM_UNCHECKED = '☐'

_WINGDINGS2_MAP = {
    '00A3': _SYM_UNCHECKED,  # Wingdings 2 0xA3 (raw codepoint)
    'F0A3': _SYM_UNCHECKED,  # Wingdings 2 0xA3 (PUA remapped)
    'F052': _SYM_CHECKED,    # Wingdings 2 0x52 (PUA remapped)
    'F051': _SYM_CHECKED,    # Wingdings 2 0x51 (alternative PUA mapping)
}

_WINGDINGS_MAP = {
    'F0FD': _SYM_UNCHECKED,  # Wingdings 0xFD
    'F0FE': _SYM_CHECKED,    # Wingdings 0xFE
}

# ── Section heading keywords ────────────────────────────────────────────────
SECTION_HEADINGS = [
    '课程考核及成绩评定',
    '课程考核与成绩评定',
    '考核及成绩评定',
    '考核与成绩评定',
    '成绩评定',
    '课程教材',
    '教材及参考资料',
    '教材与参考资料',
    '参考教材及网站',
    '参考教材与网站',
    '课程目标',
    '教学目标',
    '专业类课程的课程目标及支撑专业的毕业要求及其指标点',
]

# Labels that map directly to simple single-value fields (label: value pattern).
SIMPLE_FIELD_LABELS = {
    '课程名称': 'course_name',
    '英文名称': 'english_name',
    '课程编号': 'course_code',
    '课程性质': 'course_nature',
    '课程类型': 'course_type',
    '开课单位': 'department',
    '开课院系': 'department',
    '授课单位': 'department',
    '学时': 'total_hours',
    '总学时': 'total_hours',
    '学时数': 'total_hours',
    '适用专业': 'applicable_major',
    '先修课程': 'prerequisites',
    '大纲执笔人': 'syllabus_author',
    '执笔人': 'syllabus_author',
    '大纲更新时间': 'syllabus_update_time',
    '开课学期': 'teaching_semester',
    '大纲审核人': 'syllabus_reviewer',
    '理论': 'theory_hours',
    '实验': 'lab_hours',
    '上机': 'computer_hours',
    '其他实践环节': 'other_practice_hours',
    '分析教学': 'analysis_teaching',
    '授课班级': 'teaching_class',
    '修课人数': 'student_count',
    '课序号': 'course_seq',
    '学分': 'credits',
    '授课教师': 'teacher',
}


# ══════════════════════════════════════════════════════════════════════════════
# Low-level cell text extraction (preserves checkbox / symbol state)
# ══════════════════════════════════════════════════════════════════════════════

def _extract_cell_text(tc):
    """Extract cell text including symbol characters (Wingdings checkboxes).

    Walks ``w:r``, ``w:sym``, and ``w:ffData/w:checkBox`` elements
    so that ☑/☐ states are preserved.
    """
    paras = tc.findall('.//' + qn('w:p'))
    lines = []
    for p in paras:
        line_parts = []

        # Wingdings symbol checkboxes (Insert > Symbol)
        for r_elem in p.findall(qn('w:r')):
            for sym in r_elem.findall(qn('w:sym')):
                font = sym.get(qn('w:font'), '')
                char = sym.get(qn('w:char'), '')
                if 'Wingdings' in font:
                    mapped = _WINGDINGS2_MAP.get(char) or _WINGDINGS_MAP.get(char)
                    if mapped:
                        line_parts.append(mapped)
            for t_elem in r_elem.findall(qn('w:t')):
                if t_elem.text:
                    line_parts.append(t_elem.text)

        # Legacy form field checkboxes (Developer > Legacy Tools)
        for ffdata in p.findall('.//' + qn('w:ffData')):
            cb = ffdata.find(qn('w:checkBox'))
            if cb is not None:
                checked = cb.find(qn('w:checked'))
                default = cb.find(qn('w:default'))
                if (checked is not None and checked.get(qn('w:val')) == '1') or \
                   (default is not None and default.get(qn('w:val')) == '1'):
                    line_parts.append(_SYM_CHECKED)
                else:
                    line_parts.append(_SYM_UNCHECKED)

        lines.append(''.join(line_parts))
    return '\n'.join(lines).strip().replace('\n', ' ')


# ══════════════════════════════════════════════════════════════════════════════
# Merged-cell-aware table parsing
# ══════════════════════════════════════════════════════════════════════════════

def _get_merged_cell_map(table):
    """Parse a Word table handling gridSpan / vMerge / hMerge.

    Returns (merged_map, origins) where:
      merged_map: (row, grid_col) → cell text (with checkboxes)
      origins:    (row, grid_col) → (source_row, source_grid_col)
    """
    merged_map = {}
    origins = {}
    num_rows = len(table.rows)

    for row_idx, row in enumerate(table.rows):
        tr = row._tr
        tc_elements = tr.findall(qn('w:tc'))
        grid_col = 0

        for tc in tc_elements:
            tcpr = tc.find(qn('w:tcPr'))

            # gridSpan
            gs = 1
            if tcpr is not None:
                gs_el = tcpr.find(qn('w:gridSpan'))
                if gs_el is not None:
                    gs = int(gs_el.get(qn('w:val')))

            # vMerge / hMerge
            vm_val = None
            hm_val = None
            if tcpr is not None:
                vm_el = tcpr.find(qn('w:vMerge'))
                if vm_el is not None:
                    vm_val = vm_el.get(qn('w:val'), 'continue')
                hm_el = tcpr.find(qn('w:hMerge'))
                if hm_el is not None:
                    hm_val = hm_el.get(qn('w:val'), 'continue')

            cell_text = _extract_cell_text(tc)
            origin = (row_idx, grid_col)

            if vm_val == 'continue':
                grid_col += gs
                continue

            # Vertical merge — fill all rows
            if vm_val == 'restart':
                r = row_idx
                while r < num_rows:
                    for c in range(grid_col, grid_col + gs):
                        merged_map[(r, c)] = cell_text
                        origins[(r, c)] = origin
                    r += 1
                    if r >= num_rows:
                        break
                    next_tr = table.rows[r]._tr
                    next_tcs = next_tr.findall(qn('w:tc'))
                    g = 0
                    found_continue = False
                    for next_tc in next_tcs:
                        n_tcpr = next_tc.find(qn('w:tcPr'))
                        n_gs = 1
                        if n_tcpr is not None:
                            n_gs_el = n_tcpr.find(qn('w:gridSpan'))
                            if n_gs_el is not None:
                                n_gs = int(n_gs_el.get(qn('w:val')))
                        if g == grid_col:
                            n_vm_el = n_tcpr.find(qn('w:vMerge')) if n_tcpr is not None else None
                            if n_vm_el is not None:
                                found_continue = True
                            break
                        g += n_gs
                    if not found_continue:
                        break

            # Horizontal merge — fill all columns
            elif hm_val == 'restart':
                for c in range(grid_col, grid_col + gs):
                    merged_map[(row_idx, c)] = cell_text
                    origins[(row_idx, c)] = origin

            # Normal cell
            else:
                for c in range(grid_col, grid_col + gs):
                    merged_map[(row_idx, c)] = cell_text
                    origins[(row_idx, c)] = origin

            grid_col += gs

    return merged_map, origins


def _collect_row_cells(merged_map, origins, row_idx, max_grid_col):
    """Collect distinct logical cells in a row, ordered by grid column."""
    seen_origins = set()
    cells = []
    for col_idx in range(max_grid_col + 1):
        text = merged_map.get((row_idx, col_idx))
        if not text:
            continue
        origin = origins.get((row_idx, col_idx))
        if origin not in seen_origins:
            seen_origins.add(origin)
            cells.append((col_idx, text))
    return cells


def _build_table_grid(merged_map, origins, num_rows, num_cols):
    """Build a grid representation of a table with cell span information.

    Returns a list of rows, each a list where:
      - {'text': ..., 'colspan': N, 'rowspan': M} marks a cell that starts here
      - None marks a position covered by a spanning cell
    """
    grid = [[None] * num_cols for _ in range(num_rows)]

    origin_cells = {}
    for (r, c), origin in origins.items():
        origin_cells.setdefault(origin, []).append((r, c))

    for origin, positions in origin_cells.items():
        src_r, src_c = origin
        rows = [p[0] for p in positions]
        cols = [p[1] for p in positions]
        rowspan = max(rows) - min(rows) + 1
        colspan = max(cols) - min(cols) + 1
        text = merged_map.get((src_r, src_c), '')
        grid[src_r][src_c] = {'text': text, 'colspan': colspan, 'rowspan': rowspan}

    return grid


def _extract_all_table_kv(tables_rich):
    """Extract ALL key-value pairs from every table using the rich cell data.

    Pairs adjacent cells within each row (key → value). Also handles
    vMerge rows where a header spans two rows and sub-keys sit beside values.
    """
    all_kv = {}
    for t_idx, t in enumerate(tables_rich):
        merged_map = t['merged_map']
        origins = t['origins']
        num_rows = t['num_rows']

        max_grid_col = 0
        for (_r, c) in merged_map:
            if c > max_grid_col:
                max_grid_col = c

        table_kv = {}
        skip_rows = set()

        for row_idx in range(num_rows):
            if row_idx in skip_rows:
                continue

            row_cells = _collect_row_cells(merged_map, origins, row_idx, max_grid_col)

            # Detect vMerge: if next row's first-cell origin is the current row
            next_row = row_idx + 1
            if next_row < num_rows:
                header_row = None
                for col_idx in range(max_grid_col + 1):
                    origin = origins.get((next_row, col_idx))
                    if origin is not None:
                        if origin[0] < next_row:
                            header_row = origin[0]
                        break

                if header_row == row_idx:
                    next_cells = _collect_row_cells(merged_map, origins, next_row, max_grid_col)
                    main_key = next_cells[0][1] if next_cells else ''
                    # Remove shared first cell from the sub-row
                    if row_cells and next_cells and row_cells[0][1] == next_cells[0][1]:
                        next_cells = next_cells[1:]
                    sub_keys = row_cells[1:]
                    next_col_map = {col: text for col, text in next_cells}

                    for sk_col, sk_text in sub_keys:
                        value = next_col_map.get(sk_col, '')
                        if value:
                            full_key = f'{main_key}{sk_text}' if main_key else sk_text
                            table_kv[full_key] = value

                    skip_rows.add(next_row)
                    continue

            # Normal row: pair adjacent cells as key → value
            texts = [t for _, t in row_cells]
            for i in range(0, len(texts) - 1, 2):
                key = texts[i]
                value = texts[i + 1]
                if key:
                    table_kv[key] = value

        all_kv.update(table_kv)
    return all_kv


# ══════════════════════════════════════════════════════════════════════════════
# Utilities
# ══════════════════════════════════════════════════════════════════════════════

def _normalize(text):
    return re.sub(r'\s+', '', text)


def _label_match(text, label):
    """Check if *text* matches *label*, ignoring trailing colons."""
    t = _normalize(text).rstrip('：:').rstrip('：:')
    l = _normalize(label).rstrip('：:').rstrip('：:')
    if not t or not l:
        return False
    return t == l or t.startswith(l) or l in t


TABLE_GUIDE_PATTERNS = [
    re.compile(r'如下表[所示]*'),
    re.compile(r'下表[所示]*'),
    re.compile(r'对应关系[表图]*'),
    re.compile(r'如下[所示]*[：:]'),
    re.compile(r'[表图]如下'),
]

def _is_table_guide(text):
    """Check if a paragraph is a table/figure guide caption that should be excluded."""
    clean = _normalize(text)
    if not clean:
        return False
    for pat in TABLE_GUIDE_PATTERNS:
        if pat.search(clean):
            return True
    return False


def _is_heading(para_text, targets):
    t = _normalize(para_text)
    cleaned = re.sub(r'^[（(][一二三四五六七八九十\d]+[）)]', '', t)
    cleaned = re.sub(r'^[一二三四五六七八九十]+[、.]', '', cleaned)
    cleaned = re.sub(r'^\d+[、.)\s]+', '', cleaned)
    # Also strip leading Chinese-numeral prefixes like "五、" → ""
    cleaned = re.sub(r'^[一二三四五六七八九十]+[、.]?\s*', '', cleaned)
    cleaned = cleaned.strip()
    if not cleaned:
        return None
    best_ratio = 0.0
    best_target = None
    for target in targets:
        norm_target = _normalize(target)
        if cleaned == norm_target:
            return target
        ratio = SequenceMatcher(None, cleaned, norm_target).ratio()
        if ratio > best_ratio:
            best_ratio = ratio
            best_target = target
    return best_target if best_ratio >= 0.7 else None


def _find_section_range(paragraphs, start_heading, end_headings):
    start_targets = [start_heading] if isinstance(start_heading, str) else start_heading
    start_idx = None
    for i, p in enumerate(paragraphs):
        if _is_heading(p['text'], start_targets):
            start_idx = i
            break
    if start_idx is None:
        return None, None

    end_idx = len(paragraphs)
    for i in range(start_idx + 1, len(paragraphs)):
        if _is_heading(paragraphs[i]['text'], end_headings):
            end_idx = i
            break
    return start_idx, end_idx


# ══════════════════════════════════════════════════════════════════════════════
# Public API
# ══════════════════════════════════════════════════════════════════════════════

def parse_docx(file_path):
    """Parse a .docx file and return its paragraphs and tables.

    Tables are parsed with merged-cell awareness and checkbox state preserved.
    """
    ext = os.path.splitext(file_path)[1].lower()
    if ext != '.docx':
        return {'error': '仅支持 .docx 格式在线预览，.doc 文件请先转换为 .docx'}

    doc = Document(file_path)

    # ── Paragraphs ──────────────────────────────────────────────────────
    paragraphs = []
    for idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        style = para.style.name if para.style else 'Normal'
        paragraphs.append({
            'index': idx,
            'text': text,
            'style': style,
        })

    # ── Tables (rich: merged-cell-aware with checkbox symbols) ─────────
    tables = []
    tables_rich = []
    for t_idx, table in enumerate(doc.tables):
        merged_map, origins = _get_merged_cell_map(table)

        max_grid_col = 0
        for (_r, c) in merged_map:
            if c > max_grid_col:
                max_grid_col = c

        # Build simplified row view for display / API compatibility
        rows = []
        for r_idx in range(len(table.rows)):
            cells = _collect_row_cells(merged_map, origins, r_idx, max_grid_col)
            rows.append([t for _, t in cells])

        tables.append({
            'index': t_idx,
            'headers': rows[0] if rows else [],
            'rows': rows[1:] if len(rows) > 1 else [],
        })

        # Keep rich data for KV extraction
        tables_rich.append({
            'index': t_idx,
            'merged_map': merged_map,
            'origins': origins,
            'num_rows': len(table.rows),
        })

    # ── Build body element order (interleaved paragraphs and tables) ────
    para_lookup = {p['index']: i for i, p in enumerate(paragraphs)}
    _body_elements = []
    api_body_elements = []
    p_idx = 0
    t_idx = 0
    for child in doc.element.body:
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag == 'p':
            if p_idx in para_lookup:
                para = paragraphs[para_lookup[p_idx]]
                _body_elements.append({'type': 'p', 'text': para['text']})
                api_body_elements.append({
                    'type': 'paragraph',
                    'text': para['text'],
                    'style': para['style'],
                })
            p_idx += 1
        elif tag == 'tbl':
            if t_idx < len(tables_rich):
                rich = tables_rich[t_idx]
                max_grid_col = max((c for (_r, c) in rich['merged_map']), default=-1)
                num_cols = max_grid_col + 1
                grid = _build_table_grid(rich['merged_map'], rich['origins'], rich['num_rows'], num_cols)
                _body_elements.append({'type': 't', 'num_cols': num_cols, 'grid': grid})
                api_body_elements.append({'type': 'table', 'table_index': t_idx})
            t_idx += 1

    return {
        'paragraphs': paragraphs,
        'tables': tables,
        'body_elements': api_body_elements,
        '_tables_rich': tables_rich,
        '_body_elements': _body_elements,
    }


def _extract_section_blocks(body_elements, start_headings, end_headings):
    """Extract content blocks (paragraphs + tables) within a section boundary.

    Returns a list of blocks:
      {'type': 'paragraph', 'text': ...}
      {'type': 'table', 'num_cols': N, 'grid': [[cell|None, ...], ...]}
        where cell is {'text': ..., 'colspan': N, 'rowspan': M}
    """
    start_targets = [start_headings] if isinstance(start_headings, str) else start_headings
    start_be_idx = None
    for i, be in enumerate(body_elements):
        if be['type'] == 'p' and _is_heading(be['text'], start_targets):
            start_be_idx = i
            break
    if start_be_idx is None:
        return []

    end_be_idx = len(body_elements)
    for i in range(start_be_idx + 1, len(body_elements)):
        be = body_elements[i]
        if be['type'] == 'p' and _is_heading(be['text'], end_headings):
            end_be_idx = i
            break

    blocks = []
    for i in range(start_be_idx + 1, end_be_idx):
        be = body_elements[i]
        if be['type'] == 'p':
            if be['text']:
                blocks.append({'type': 'paragraph', 'text': be['text']})
        elif be['type'] == 't':
            blocks.append({
                'type': 'table',
                'num_cols': be['num_cols'],
                'grid': be['grid'],
            })
    return blocks


def _extract_evaluation_table(evaluation_blocks):
    """Extract evaluation items and course-objective matrix from the eval table.

    Finds the table containing '评价依据及成绩比例(%)' and returns:
        eval_items:     [{'name': '课堂表现', 'percentage': '10%'}, ...]
        course_obj_items: ['目标1：...', '目标2：...', ...]
        matrix:         [[value, ...], ...]  # course_obj × eval_item
        course_obj_count: int
    """
    if not evaluation_blocks or not isinstance(evaluation_blocks, list):
        return None

    for block in evaluation_blocks:
        if block.get('type') != 'table':
            continue
        grid = block.get('grid')
        if not grid or len(grid) < 3:
            continue

        num_cols = block.get('num_cols', len(grid[0]) if grid else 0)

        # ── Locate the evaluation header row and column ranges ──────────
        eval_header_row = None
        eval_col_start = None
        eval_col_end = None
        obj_col_idx = None

        for r_idx in range(min(3, len(grid))):
            row = grid[r_idx]
            for c in range(num_cols):
                cell = row[c] if c < len(row) else None
                if cell is None:
                    continue
                text = str(cell.get('text', ''))
                if '评价依据及成绩比例' in text:
                    eval_header_row = r_idx
                    eval_col_start = c
                    eval_col_end = c + cell.get('colspan', 1) - 1
                    break
                if '课程目标' in text and obj_col_idx is None:
                    obj_col_idx = c
            if eval_header_row is not None:
                break

        if eval_header_row is None or eval_col_start is None or obj_col_idx is None:
            continue

        # ── Read evaluation sub-item names from the row below header ─────
        sub_row_idx = eval_header_row + 1
        if sub_row_idx >= len(grid):
            continue

        eval_items = []
        for c in range(eval_col_start, eval_col_end + 1):
            cell = grid[sub_row_idx][c] if c < len(grid[sub_row_idx]) else None
            if cell is None:
                continue
            text = str(cell.get('text', '')).replace('\n', '').replace('\r', '').strip()
            # Extract the designed percentage from cell text (e.g. "10%" from "课堂表现 10%")
            pct_match = re.search(r'(\d+(?:\.\d+)?)%$', text)
            cell_pct = pct_match.group(1) + '%' if pct_match else None
            # Strip trailing percentage from the name — the designed percentage is stored separately
            text = re.sub(r'\s*\d+(?:\.\d+)?%$', '', text).strip()
            if text:
                eval_items.append({'name': text, 'percentage': cell_pct})

        if not eval_items:
            continue

        # ── Read course objective rows (data rows after sub-header) ─────
        course_obj_items = []
        matrix = []
        data_start = sub_row_idx + 1

        for r_idx in range(data_start, len(grid)):
            row = grid[r_idx]
            obj_cell = row[obj_col_idx] if obj_col_idx < len(row) else None
            if obj_cell is None:
                continue
            obj_text = str(obj_cell.get('text', '')).replace('\n', '').replace('\r', '').strip()
            if not obj_text:
                continue
            # Skip 合计/总计 row
            if '合计' in obj_text or '总计' in obj_text:
                continue

            course_obj_items.append(obj_text)
            row_values = []
            for c in range(eval_col_start, eval_col_end + 1):
                cell = row[c] if c < len(row) else None
                val = str(cell.get('text', '')).replace('\n', '').replace('\r', '').strip() if cell else ''
                row_values.append(val)
            matrix.append(row_values)

        if not course_obj_items:
            continue

        # ── Calculate total percentage for each eval item ────────────────
        # Sum values in each column, parse as numbers
        for ci in range(len(eval_items)):
            # Only compute percentage from data rows if the cell text didn't include one
            if eval_items[ci]['percentage'] is not None:
                continue
            total_pct = 0
            for ri in range(len(matrix)):
                try:
                    total_pct += float(matrix[ri][ci]) if matrix[ri][ci] else 0
                except (ValueError, IndexError):
                    pass
            eval_items[ci]['percentage'] = f'{int(total_pct)}%' if total_pct == int(total_pct) else f'{total_pct:.1f}%'

        return {
            'eval_items': eval_items,
            'course_obj_items': course_obj_items,
            'matrix': matrix,
            'course_obj_count': len(course_obj_items),
        }

    return None


def extract_syllabus_fields(paragraphs, tables, tables_rich=None, body_elements=None):
    """Extract structured syllabus fields from parsed .docx content.

    Combines:
    1. Rich table KV extraction (ALL fields including checkbox states)
    2. Labeled paragraph extraction (fallback for non-table fields)
    3. Section-based content extraction (课程目标, 教材, 评价标准)

    Returns a dict with keys matching the report template fields.
    """
    # 1. Rich table extraction — capture ALL key-value pairs including checkboxes
    fields = {}
    if tables_rich:
        all_kv = _extract_all_table_kv(tables_rich)
        # Map known labels to canonical field names
        for label, key in SIMPLE_FIELD_LABELS.items():
            for kv_key, value in all_kv.items():
                if _label_match(kv_key, label):
                    if key not in fields and value.strip():
                        fields[key] = value.strip()
                        break

        # Store ALL raw table KV pairs for report generation use
        fields['_all_table_kv'] = all_kv

    # 2. Simple fields from labeled paragraphs (fill gaps not found in tables)
    for p in paragraphs:
        text = p['text'].strip()
        for label, key in SIMPLE_FIELD_LABELS.items():
            if key in fields:
                continue
            # Match "label：value" or "label: value" patterns
            m = re.match(r'(.+?)\s*[：:]\s*(.+)', text)
            if m:
                maybe_label = m.group(1).strip()
                val = m.group(2).strip()
                if _label_match(maybe_label, label) and val:
                    fields[key] = val
                    break

    # 3. 课程目标
    start, end = _find_section_range(
        paragraphs, ['课程目标', '教学目标'], SECTION_HEADINGS)
    if start is not None:
        texts = []
        for i in range(start + 1, end):
            t = paragraphs[i]['text'].strip()
            if t:
                texts.append(t)
        while texts and _is_table_guide(texts[-1]):
            texts.pop()
        if texts:
            fields['course_objectives'] = '\n'.join(texts)

        # ── Extract course objective count from table ──────────────────
        if body_elements:
            obj_blocks = _extract_section_blocks(
                body_elements, ['课程目标', '教学目标'], SECTION_HEADINGS)
            for block in obj_blocks:
                if block.get('type') != 'table':
                    continue
                grid = block.get('grid')
                if not grid:
                    continue
                # Find the column with '课程目标' header
                obj_col = None
                seq_col = None
                for r_idx in range(min(3, len(grid))):
                    row = grid[r_idx]
                    for c in range(len(row)):
                        cell = row[c]
                        if cell is None:
                            continue
                        text = str(cell.get('text', ''))
                        if '课程目标' in text or '课程目标' in text:
                            obj_col = c
                        if '序号' in text and seq_col is None:
                            seq_col = c
                if obj_col is not None:
                    count = 0
                    max_seq = 0
                    for r_idx in range(2, len(grid)):
                        row = grid[r_idx]
                        cell = row[obj_col] if obj_col < len(row) else None
                        if cell is None:
                            continue
                        text = str(cell.get('text', '')).strip()
                        if not text or '合计' in text or '总计' in text:
                            continue
                        count += 1
                        if seq_col is not None and seq_col < len(row):
                            seq_cell = row[seq_col]
                            if seq_cell:
                                try:
                                    max_seq = max(max_seq, int(str(seq_cell.get('text', '0'))))
                                except ValueError:
                                    pass
                    # Cross-validate
                    if count > 0:
                        fields['course_obj_count'] = max(count, max_seq)
                    break

    # 4. 选用教材
    textbook_start, textbook_end = _find_section_range(
        paragraphs, '课程教材', SECTION_HEADINGS)
    if textbook_start is not None:
        ref_start, _ = _find_section_range(
            paragraphs, ['参考教材及网站', '参考教材与网站'], [])
        effective_end = ref_start if ref_start is not None else textbook_end
        texts = []
        for i in range(textbook_start + 1, effective_end):
            t = paragraphs[i]['text'].strip()
            if t:
                texts.append(t)
        if texts:
            fields['textbook'] = '\n'.join(texts)

    # 5. 课程考核及成绩评定 — extract as blocks (paragraphs + tables)
    eval_headings = ['课程考核及成绩评定', '课程考核与成绩评定', '考核及成绩评定', '考核与成绩评定']
    if body_elements:
        blocks = _extract_section_blocks(
            body_elements, eval_headings, SECTION_HEADINGS)
        if blocks:
            fields['evaluation_standards'] = blocks
            # Build the evaluation dataframe
            eval_table = _extract_evaluation_table(blocks)
            if eval_table:
                fields['eval_items'] = eval_table['eval_items']
                fields['course_obj_items'] = eval_table['course_obj_items']
                fields['eval_matrix'] = eval_table['matrix']
                if 'course_obj_count' not in fields:
                    fields['course_obj_count'] = eval_table['course_obj_count']
    else:
        eval_start, eval_end = _find_section_range(paragraphs, eval_headings, SECTION_HEADINGS)
        if eval_start is not None:
            texts = []
            for i in range(eval_start + 1, eval_end):
                t = paragraphs[i]['text'].strip()
                if t:
                    texts.append(t)
            if texts:
                fields['evaluation_standards'] = '\n'.join(texts)

    # 6. Fill defaults
    all_keys = (
        list(SIMPLE_FIELD_LABELS.values())
        + ['course_objectives', 'textbook', 'evaluation_standards',
           'course_obj_count', 'eval_items', 'course_obj_items', 'eval_matrix']
    )
    for k in all_keys:
        if k not in fields:
            fields[k] = '未填写'

    return fields
